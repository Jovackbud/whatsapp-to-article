import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from io import BytesIO
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors  # Added import for color support
from datetime import datetime
import re
import logging # New import for logging errors
from typing import Optional, Tuple, List # New import for type hints
from datetime import datetime
import re
import pyperclip  # Added import for pyperclip

logging.basicConfig(level=logging.INFO)

class DocumentExporter:
    """
    Handles the export of processed content into various document formats like Word and PDF.
    Focuses on content formatting and document generation, separating concerns from Streamlit UI.
    """
    
    @staticmethod
    def create_word_doc(content: str, title: Optional[str] = None, author: str = "WhatsApp Chat Converter") -> bytes:
        """Create a Word document from the content."""
        try:
            doc = Document()
            
            # Set document title
            if title:
                title_paragraph = doc.add_heading(title, 0) # Level 0 is 'Title' style
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            core_props = doc.core_properties
            core_props.author = author
            core_props.created = datetime.now()
            
            # Process content into blocks
            processed_blocks = DocumentExporter._process_content_lines(content)
            
            for block_type, block_text in processed_blocks:
                if block_type == 'paragraph':
                    doc.add_paragraph(block_text)
                elif block_type == 'heading':
                    heading_info = DocumentExporter._parse_heading(block_text)
                    if heading_info:
                        heading_level, clean_heading = heading_info
                        # docx heading levels are 0 for title, 1 for Heading 1, 2 for Heading 2 etc.
                        # So, if MD is #, level is 1. If MD is ##, level is 2.
                        doc.add_heading(clean_heading, level=heading_level)
                    else:
                        # Fallback if _is_heading returned true but _parse_heading failed (should not happen with strict _is_heading)
                        doc.add_heading(block_text, level=2) # Default to Heading 2
            
            # Save to BytesIO
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer.getvalue()
            
        except Exception as e:
            logging.exception(f"Error creating Word document: {e}")
            raise RuntimeError(f"Failed to create Word document: {e}. Please try again or check the content format.")
    
    @staticmethod
    def create_pdf(content: str, title: Optional[str] = None, author: str = "WhatsApp Chat Converter") -> bytes:
        """Create a PDF document from the content."""
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Get ReportLab's default styles
            styles = getSampleStyleSheet()
            
            # Custom styles for various heading levels and body
            h1_style = ParagraphStyle(
                'ArticleHeading1',
                parent=styles['h1'], # Use reportlab's h1 as base
                fontSize=18,
                alignment=TA_CENTER,
                spaceAfter=24,
                textColor=colors.HexColor('#2E4053') # Dark blue/grey
            )
            
            h2_style = ParagraphStyle(
                'ArticleHeading2',
                parent=styles['h2'], # Use reportlab's h2 as base
                fontSize=14,
                spaceBefore=16,
                spaceAfter=8,
                textColor=colors.HexColor('#34495E') # Slightly lighter blue/grey
            )

            h3_style = ParagraphStyle(
                'ArticleHeading3',
                parent=styles['h3'], # Use reportlab's h3 as base
                fontSize=12,
                spaceBefore=12,
                spaceAfter=6,
                textColor=colors.HexColor('#4A637F') # Even lighter blue/grey
            )
            
            body_style = ParagraphStyle(
                'ArticleBody',
                parent=styles['Normal'],
                fontSize=11,
                alignment=TA_JUSTIFY,
                spaceBefore=6,
                spaceAfter=6,
                leftIndent=0,
                rightIndent=0
            )
            
            story: List = []
            
            # Add title if provided (will be treated as a main heading)
            if title:
                story.append(Paragraph(title, h1_style))
                story.append(Spacer(1, 12))
            
            # Process content into blocks
            processed_blocks = DocumentExporter._process_content_lines(content)
            
            for block_type, block_text in processed_blocks:
                if block_type == 'paragraph':
                    story.append(Paragraph(block_text, body_style))
                elif block_type == 'heading':
                    heading_info = DocumentExporter._parse_heading(block_text)
                    if heading_info:
                        heading_level, clean_heading = heading_info
                        if heading_level == 1:
                            story.append(Paragraph(clean_heading, h1_style))
                            story.append(Spacer(1, 12))
                        elif heading_level == 2:
                            story.append(Paragraph(clean_heading, h2_style))
                            story.append(Spacer(1, 8))
                        elif heading_level == 3:
                            story.append(Paragraph(clean_heading, h3_style))
                            story.append(Spacer(1, 6))
                        else: # Fallback for higher levels or unexpected
                            story.append(Paragraph(clean_heading, body_style))
                            story.append(Spacer(1, 6))
                    else:
                        # Fallback if _is_heading returned true but _parse_heading failed
                        story.append(Paragraph(block_text, body_style))
                        story.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            
            return buffer.getvalue()
            
        except Exception as e:
            logging.exception(f"Error creating PDF document: {e}")
            raise RuntimeError(f"Failed to create PDF document: {e}. Please try again or check the content format.")

    @staticmethod
    def _is_heading(line: str) -> bool:
        """
        Determine if a line should be treated as a heading based on Markdown syntax.
        Returns True if it's a heading, False otherwise.
        """
        stripped_line = line.strip()
        # Explicit Markdown heading markers as per prompt instructions
        return stripped_line.startswith('# ') or \
               stripped_line.startswith('## ') or \
               stripped_line.startswith('### ')

    @staticmethod
    def _parse_heading(line: str) -> Optional[Tuple[int, str]]:
        """
        Parses a heading line and returns its level and cleaned text.
        Returns (level, cleaned_text) or None if not a recognized heading format.
        """
        stripped_line = line.strip()
        if stripped_line.startswith('# '):
            return 1, stripped_line[2:].strip()
        elif stripped_line.startswith('## '):
            return 2, stripped_line[3:].strip()
        elif stripped_line.startswith('### '):
            return 3, stripped_line[4:].strip()
        
        return None
    
    @staticmethod
    def _process_content_lines(content: str) -> List[Tuple[str, str]]:
        """
        Processes content lines and categorizes them as 'paragraph' or 'heading'.
        Returns a list of tuples: [('type', 'text'), ...]
        """
        lines = content.split('\n')
        processed_blocks = []
        current_paragraph_lines = []

        def add_current_paragraph():

            if current_paragraph_lines:
                processed_blocks.append(('paragraph', " ".join(current_paragraph_lines)))
                current_paragraph_lines.clear()

        for line in lines:
            line_stripped = line.strip()
            
            if not line_stripped: # Empty line, signifies paragraph break
                add_current_paragraph()
                continue
            
            if DocumentExporter._is_heading(line_stripped):
                add_current_paragraph() # Add any accumulated paragraph before the heading
                processed_blocks.append(('heading', line_stripped))
            else:
                # Accumulate paragraph content
                current_paragraph_lines.append(line_stripped)
        
        add_current_paragraph() # Add any remaining paragraph
        return processed_blocks


    @staticmethod
    def generate_filename(title: Optional[str] = None, file_type: str = "docx") -> str:
        """Generate a filename for download"""
        if title:
            # Clean title for filename
            clean_title = re.sub(r'[^\w\s-]', '', title)
            clean_title = re.sub(r'[-\s]+', '-', clean_title)
            filename = clean_title + "." + file_type
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = "converted_article_" + timestamp + "." + file_type
        
        return filename
